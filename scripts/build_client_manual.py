from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "client-manual"
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "Swift-Financial-Client-User-Guide.docx"

INDIGO = "3730A3"
INDIGO_DARK = "312E81"
INDIGO_LIGHT = "EEF2FF"
INK = "1F2937"
MUTED = "6B7280"
GRAY = "E5E7EB"
GREEN = "166534"
AMBER = "92400E"
RED = "991B1B"
WHITE = "FFFFFF"

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def set_font(run, name="Aptos", size=None, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.15
for name, size, color, before, after in [
    ("Title", 29, WHITE, 0, 10),
    ("Heading 1", 17, INDIGO_DARK, 16, 7),
    ("Heading 2", 13, INDIGO, 11, 5),
    ("Heading 3", 11, INK, 8, 3),
]:
    s = styles[name]
    s.font.name = "Aptos Display" if name != "Normal" else "Aptos"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

for list_style in ("List Bullet", "List Number"):
    s = styles[list_style]
    s.font.name = "Aptos"
    s.font.size = Pt(10.5)
    s.paragraph_format.left_indent = Inches(0.38)
    s.paragraph_format.first_line_indent = Inches(-0.19)
    s.paragraph_format.space_after = Pt(3)
    s.paragraph_format.line_spacing = 1.15

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)

header = section.header
hp = header.paragraphs[0]
hp.text = "SWIFT FINANCIAL  |  CLIENT USER GUIDE"
set_font(hp.runs[0], size=8, bold=True, color=MUTED)
hp.paragraph_format.space_after = Pt(0)
footer = section.footer
fp = footer.paragraphs[0]
fp.add_run("Client copy  |  Version 1.0  |  September 2026")
set_font(fp.runs[0], size=8, color=MUTED)
add_page_number(fp)

def page_break():
    doc.add_page_break()

def p(text="", style=None, bold_prefix=None):
    para = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = para.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = para.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = para.add_run(text)
        set_font(r)
    return para

def heading(text, level=1):
    return doc.add_heading(text, level=level)

def bullet(text): return p(text, "List Bullet")
def number(text): return p(text, "List Number")

def callout(title, text, kind="info"):
    colors = {"info": (INDIGO_LIGHT, INDIGO_DARK), "warning": ("FEF3C7", AMBER), "danger": ("FEE2E2", RED), "success": ("DCFCE7", GREEN)}
    fill, color = colors[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.72)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 130, 160, 130, 160)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(2)
    r = para.add_run(title.upper())
    set_font(r, size=9, bold=True, color=color)
    para2 = cell.add_paragraph()
    para2.paragraph_format.space_after = Pt(0)
    r2 = para2.add_run(text)
    set_font(r2, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.style = "Table Grid"
    if widths is None:
        widths = [6.72 / len(headers)] * len(headers)
    for idx, (cell, label, width) in enumerate(zip(t.rows[0].cells, headers, widths)):
        cell.width = Inches(width)
        set_cell_shading(cell, INDIGO_DARK)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(label)
        set_font(run, size=9, bold=True, color=WHITE)
    set_repeat_table_header(t.rows[0])
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for cell, value, width in zip(cells, row, widths):
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if ri % 2: set_cell_shading(cell, "F9FAFB")
            run = cell.paragraphs[0].add_run(str(value))
            set_font(run, size=9.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t

# Cover
cover = doc.add_table(rows=1, cols=1)
cover.alignment = WD_TABLE_ALIGNMENT.CENTER
cover.autofit = False
cover.columns[0].width = Inches(6.72)
cell = cover.cell(0, 0)
set_cell_shading(cell, INDIGO_DARK)
set_cell_margins(cell, 620, 420, 620, 420)
brand = cell.paragraphs[0]
brand.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = brand.add_run("SWIFT FINANCIAL")
set_font(r, size=12, bold=True, color="C7D2FE")
title = cell.add_paragraph()
title.paragraph_format.space_before = Pt(26)
title.paragraph_format.space_after = Pt(12)
r = title.add_run("Client User Guide")
set_font(r, "Aptos Display", 30, True, WHITE)
sub = cell.add_paragraph()
r = sub.add_run("A practical guide to secure, accurate day-to-day use")
set_font(r, size=14, color="E0E7FF")
meta = cell.add_paragraph()
meta.paragraph_format.space_before = Pt(36)
r = meta.add_run("Version 1.0  |  September 2026")
set_font(r, size=10, color="C7D2FE")
p("Prepared for: [Client organisation]")
p("System URL: [Production URL]")
p("Support contact: [Name, email and telephone]")
callout("Purpose", "This guide is for authorised client staff using Swift Financial. Available menus and actions depend on the permissions assigned to each user.")

page_break()
heading("Document control", 1)
table(["Item", "Details"], [
    ("Document owner", "[Client system owner]"),
    ("Application", "Swift Financial - Front Office Service Activities and related modules"),
    ("Version", "1.0"),
    ("Issued", "September 2026"),
    ("Audience", "Tellers, operations staff, supervisors, finance staff, administrators and authorised reviewers"),
    ("Review cycle", "Review after each major release, workflow change or permission-model change"),
], [1.75, 4.97])
heading("How to use this guide", 2)
bullet("Start with Quick start if you are a new user.")
bullet("Use the module overview to identify where a task belongs.")
bullet("Follow the numbered workflow and complete the checks before submitting.")
bullet("Ask a supervisor before repeating, reversing or correcting a financial transaction.")
callout("Scope note", "This manual reflects the application routes and workflows present in the September 2026 client build. Your organisation may hide modules that are not licensed, configured or assigned to your role.", "info")

heading("Contents", 1)
for item in [
    "1. Quick start", "2. Account security", "3. Navigating the application", "4. Common screen patterns",
    "5. Module overview", "6. Registry and customer records", "7. Accounts and product setup",
    "8. Front Office operations", "9. Loans and check-off", "10. Batch procedures and approvals",
    "11. Reports, statements and messaging", "12. Administration", "13. Troubleshooting",
    "14. Daily operating checklists", "15. Support and handover details"
]: p(item)

page_break()
heading("1. Quick start", 1)
heading("Sign in", 2)
number("Open the production URL supplied by your system administrator.")
number("Enter your email address or username and password.")
number("Select Login. Wait for the home page and your permitted navigation areas to load.")
number("If prompted on first sign-in, enter the temporary password and choose a new password.")
callout("Password rules", "A new password must contain at least 6 characters, including an uppercase letter, a lowercase letter, a number and a symbol.", "info")
heading("Your first five checks", 2)
table(["Check", "What to confirm"], [
    ("Identity", "The displayed user account is yours; never work under another person's login."),
    ("Workspace", "The selected module matches the task you intend to perform."),
    ("Branch/context", "The applicable branch, teller, treasury, period or account context is correct."),
    ("Permissions", "Only expected menus and actions are visible."),
    ("Connectivity", "Lists load without a persistent error message."),
], [1.45, 5.27])
heading("Safe transaction habit", 2)
p("Search -> select -> verify -> enter -> review -> submit -> record the reference.")
callout("Never double-submit", "If the screen is still processing, do not click the action again. If the outcome is uncertain, search the relevant register or transaction journal before retrying.", "warning")

page_break()
heading("2. Account security", 1)
bullet("Keep passwords private. Administrators and support staff should never ask you to disclose your password.")
bullet("Use a unique password and change temporary credentials immediately.")
bullet("Lock or sign out when leaving the workstation.")
bullet("Do not export customer or financial information unless the business purpose and storage location are approved.")
bullet("Report unexpected menus, missing permissions or suspicious activity to the system owner.")
heading("Access is role-based", 2)
p("The application builds the navigation menu from the signed-in user's roles and module permissions. A user may see fewer workspaces, folders and actions than a colleague. An Access denied screen means the route is controlled and has not been granted to the current user.")
heading("When sign-in fails", 2)
table(["Message or symptom", "Response"], [
    ("Missing fields", "Enter both username/email and password."),
    ("Login failed", "Check spelling and Caps Lock, then retry once. Contact the administrator if it continues."),
    ("Initial password change fails", "Confirm both new-password entries match and meet all password rules."),
    ("Navigation cannot load", "Select Retry. If it persists, record the time and report it."),
    ("Access denied", "Return to an authorised menu. Ask the role administrator if the task is part of your duties."),
], [2.0, 4.72])

page_break()
heading("3. Navigating the application", 1)
p("The application uses a top navigation bar, a workspace rail and a contextual sidebar. On smaller screens, use the menu button to open the navigation panel.")
table(["Area", "Purpose"], [
    ("Home", "Shows permitted workspaces and provides a starting point for common functions."),
    ("Workspace rail", "Switches between major areas such as Registry, Accounts, Front Office or Administration."),
    ("Context sidebar", "Shows folders and pages available in the active workspace."),
    ("Global search", "Locates supported screens or records where enabled."),
    ("Page header", "Identifies the current function and contains primary actions such as Add or Create."),
    ("Alerts/dialogs", "Confirm success, warn about missing information or request confirmation for a sensitive action."),
], [1.55, 5.17])
heading("Responsive use", 2)
bullet("Desktop: both navigation panels remain visible.")
bullet("Tablet: select a workspace icon to open its menu.")
bullet("Mobile: use the menu button; close the overlay after choosing a page.")

heading("4. Common screen patterns", 1)
table(["Pattern", "How to use it"], [
    ("Searchable customer picker", "Type a supported identifier or name, wait for server results, select the correct customer and verify the retained selection."),
    ("List/register", "Use search or filters, open a row for details, and use Prev/Next to move through pages."),
    ("Drawer/form", "Complete required fields, scroll through the form, review, then use the fixed submit action."),
    ("Tabs", "Tabs may represent transaction types or workflow stages. Confirm the active tab before acting."),
    ("Status badge", "Green normally indicates active/authorised, amber pending/warning, blue informational/posted, red rejected/locked and gray neutral."),
    ("Confirmation alert", "Read the full action and target. Cancel if any identifier, amount, date or status is wrong."),
], [1.65, 5.07])

page_break()
heading("5. Module overview", 1)
table(["Module", "Typical functions", "Typical users"], [
    ("Registry / Membership", "Customers, members, employers, zones, divisions, linkages and file tracking", "Member service and registry staff"),
    ("Accounts", "Customer accounts, products, charges, journals, periods, bank reconciliation and budgets", "Finance and product administrators"),
    ("Front Office", "Teller receipts/payments, treasury, transfers, cheques, deposits, closure and end of day", "Tellers and supervisors"),
    ("Loaning", "Loan cases, appraisal, approval, verification, guarantors, restructuring and check-off", "Credit and loans teams"),
    ("Command Hub", "Approval requests, messages and operational utilities", "Approvers and supervisors"),
    ("Reports", "Loan, finance, member, statutory and user-defined reports", "Managers, finance and auditors"),
    ("Administration", "Users, roles, modules, workflows, banks, branches, locations and audit logs", "System administrators"),
    ("Other configured areas", "Inventory, procurement/control, HR, payroll, fixed assets and microcredit", "Specialist teams"),
], [1.35, 3.55, 1.82])
callout("Permission-dependent", "Not every route in the installed application is necessarily enabled for production use. Follow the visible, assigned navigation rather than bookmarked or manually typed addresses.", "warning")

page_break()
heading("6. Registry and customer records", 1)
heading("Find and review a customer", 2)
number("Open Registry or Membership and choose Customers or Members.")
number("Search using the available customer identifiers or name fields.")
number("Select the matching record and verify name, identifier and organisational details.")
number("Use the detail tabs to review information, accounts, next of kin or statements where available.")
heading("Register or amend a record", 2)
number("Select the appropriate create/register action.")
number("Enter the details exactly as supported by source documentation.")
number("Choose the correct branch, employer, division, zone or station linkage.")
number("Check contact details and identifiers for transposition errors.")
number("Submit once and record the confirmation or customer reference.")
callout("Customer identity", "Do not rely on name alone. Use at least one additional identifier before viewing, changing or transacting on a customer record.", "danger")
heading("Related functions", 2)
bullet("Delegates and directors maintain authorised related persons.")
bullet("Station and branch linkage assign the appropriate organisational relationship.")
bullet("Conditional lending records product-specific conditions.")
bullet("File Tracking supports dispatch, receipt, recall and catalogue activities.")

page_break()
heading("7. Accounts and product setup", 1)
p("Accounts contains both operational inquiries and setup data. Setup changes can affect future transactions, charges and postings, so they should be restricted to authorised staff and tested under change control.")
heading("Customer account workflow", 2)
number("Find and verify the customer.")
number("Select the correct financial product and branch/context.")
number("Review account identifiers, status and linked facilities.")
number("Submit the registration or maintenance action.")
number("Confirm the new or changed account in the register before continuing.")
heading("High-impact setup areas", 2)
table(["Area", "Control expectation"], [
    ("Savings, investment and loan products", "Use approved product specifications; independently review rates, limits, terms and charge settings."),
    ("Chart of accounts and mappings", "Use the approved accounting structure; validate downstream posting behaviour."),
    ("Charges, levies and commissions", "Confirm amount/rate, bearer, recovery mode, effective use and any exemptions."),
    ("Posting periods", "Keep dates and open/closed status aligned with the finance calendar."),
    ("Tellers, treasuries and bank linkages", "Verify assignment and currency/branch context before operational use."),
], [2.25, 4.47])
callout("Irreversible processing", "Posting-period closing creates fiscal closing journals. Perform it only under an approved finance procedure, after reconciliation, backup and supervisory sign-off.", "danger")

page_break()
heading("8. Front Office operations", 1)
heading("Teller receipts and payments", 2)
p("Savings Receipts/Payments supports the main teller cycle, including deposits, withdrawals, cheque deposits and payment vouchers. The active transaction type controls the fields and processing route.")
number("Confirm you are assigned to the correct teller/treasury and operational date.")
number("Choose the transaction type.")
number("Search for and select the customer/account; verify identity and account status.")
number("Enter the amount, instrument/reference and required narrative.")
number("Review the transaction summary, especially debit/credit direction and cash versus cheque treatment.")
number("Submit once, retain the generated reference and issue the approved receipt or voucher.")
heading("Treasury and cash management", 2)
bullet("Cash Management controls teller/treasury cash movements.")
bullet("Cash Withdrawal Requests follow a create, browse, authorise or reject lifecycle.")
bullet("Fiscal Counts provide a read-only catalogue for recorded counts.")
heading("Other Front Office functions", 2)
table(["Function", "Key check"], [
    ("Transfers", "Select Cash or Cheque and verify both source and destination."),
    ("Cheques", "Confirm the active Catalogue, Bank or Clear tab and cheque status."),
    ("Fixed deposits", "Confirm product/type, term, maturity instruction and funding account."),
    ("Expense payables", "Confirm beneficiary, expense purpose, amount and approval evidence."),
    ("Account closure", "Confirm the account, balance treatment, charges and required authorisation."),
    ("In-house cheques", "Verify drawer, payee, account, amount and instrument details."),
    ("Automated clearing", "Validate source file/batch, totals, duplicates and processing result."),
], [1.7, 5.02])
callout("Cash control", "Never share tills, teller credentials or transaction references. Count cash in accordance with the organisation's dual-control and till-balancing procedure.", "danger")

page_break()
heading("9. Loans and check-off", 1)
heading("Loan-case lifecycle", 2)
p("The configured loan-case pipeline separates duties across Registration, Appraisal, Approval, Verification/Audit and, where necessary, Cancellation.")
table(["Stage", "Primary responsibility"], [
    ("Registration", "Capture the applicant, requested facility, purpose, product and supporting details."),
    ("Appraisal", "Assess eligibility, affordability, security, guarantors, income adjustments and policy compliance."),
    ("Approval", "Record the authorised decision and terms within delegated limits."),
    ("Verification / Audit", "Independently confirm evidence and approved terms before disbursement processing."),
    ("Cancellation", "Cancel only an eligible case, with a documented reason and correct authority."),
], [1.65, 5.07])
heading("Guarantors and restructuring", 2)
bullet("Guarantor Attachment supports attachment, substitution and relieving/history workflows.")
bullet("Guarantor Management adds a guarantor to an already registered case.")
bullet("Restructuring is account-based; confirm the customer account and approved revised terms.")
heading("Check-off workflow", 2)
number("Create or confirm an open data period.")
number("Capture/import entries against the correct customer product accounts.")
number("Reconcile control totals and exceptions.")
number("Close the period only after review and approval.")
number("Use Catalogue for read-only follow-up and evidence.")

page_break()
heading("10. Batch procedures and approvals", 1)
heading("Batch stages", 2)
p("Batch procedures use three controlled stages: Origination, Verification and Authorization. Each stage may include multiple batch-type tabs.")
table(["Stage", "Operator check"], [
    ("Origination", "Correct batch type, source, period, row count, control total and duplicate check."),
    ("Verification", "Independent comparison to source evidence; investigate all rejected or unmatched lines."),
    ("Authorization", "Authority, totals, exception resolution, accounting date and final release decision."),
], [1.45, 5.27])
heading("Command Hub approvals", 2)
number("Open Command Hub > Approval Requests.")
number("Filter or select the pending request.")
number("Open the detail and verify the originator, action, target record, amount and supporting context.")
number("Approve or reject according to delegated authority. Enter a useful reason when required.")
number("Confirm the resulting status and retain evidence where policy requires it.")
callout("Segregation of duties", "An originator should not verify or authorise the same transaction unless the organisation has explicitly approved that control model.", "warning")

page_break()
heading("11. Reports, statements and messaging", 1)
heading("Run a report", 2)
number("Choose the report family: Loan, Financial, SASRA, Member Statement or User-Defined Reports.")
number("Set the customer, branch, account, period and other filters as applicable.")
number("Review the parameters before generating the report.")
number("Compare headline totals with the relevant register or control account.")
number("Export only when necessary and store the file in an approved location.")
heading("Statements", 2)
p("Customer Account Statement and Member Statement functions may offer preview, mini/full statement or PDF output. Confirm the customer and account identifiers before displaying or distributing the statement.")
heading("Messaging", 2)
table(["Function", "Guidance"], [
    ("Text alerts", "Review the alert history/reference information available to your role."),
    ("Email alerts", "Use history/detail/compose functions; verify recipients and attachments before sending."),
    ("Instant messaging", "Use authorised direct or group conversations; do not place passwords or unnecessary sensitive data in messages."),
], [1.65, 5.07])
callout("Confidentiality", "A report or statement remains confidential after export. Check recipients, file permissions and retention requirements before sharing.", "danger")

page_break()
heading("12. Administration", 1)
p("Administration functions are intended for designated system administrators and control owners.")
table(["Function", "Administrator responsibility"], [
    ("Users", "Create and maintain named accounts; disable leavers promptly; avoid shared accounts."),
    ("Roles and permission types", "Apply least privilege and obtain business-owner approval for access changes."),
    ("Modules", "Maintain the navigation/permission catalogue carefully; test visibility after change."),
    ("Workflows", "Configure stages and approvals in line with delegated authority and segregation of duties."),
    ("Company, branches and locations", "Keep organisational master data current and avoid duplicate records."),
    ("Banks and bank branches", "Validate identifiers and linkages before financial use."),
    ("Audit logs", "Use for investigation and review; preserve evidence and restrict access."),
], [2.0, 4.72])
heading("Access review checklist", 2)
bullet("Review privileged roles and dormant accounts regularly.")
bullet("Reconcile joiners, movers and leavers against approved HR records.")
bullet("Test that a sample user can see required pages and cannot see restricted pages.")
bullet("Retain the access request, approval and implementation evidence.")

page_break()
heading("13. Troubleshooting", 1)
table(["Problem", "Safe response"], [
    ("A list remains empty", "Check filters and page number. Confirm you have permission and that the record exists in the selected branch/period."),
    ("A save fails", "Read the alert, preserve entered details, correct validation errors and retry once."),
    ("A request times out", "Do not immediately repeat a financial action. Search the register/journal first."),
    ("Wrong record was selected", "Cancel before submission. If already submitted, stop and follow the approved correction/reversal process."),
    ("Unexpected status", "Refresh/reopen the record and check workflow history. Escalate with the reference and timestamp."),
    ("Permission or navigation issue", "Capture the page name and required duty; ask the role administrator to review access."),
    ("Persistent system error", "Record the user, time, page, action, reference and exact message; provide a screenshot without exposing passwords."),
], [2.05, 4.67])
heading("Support ticket minimum information", 2)
bullet("Your name, username and branch - never your password.")
bullet("Date/time and application page.")
bullet("What you were trying to do and the steps immediately before the error.")
bullet("Customer/account/transaction reference, masked where appropriate.")
bullet("Exact error message and a safe screenshot.")
bullet("Whether the action appears in the relevant register or journal.")

page_break()
heading("14. Daily operating checklists", 1)
heading("Start of day", 2)
bullet("Sign in with your own account and confirm the correct permissions.")
bullet("Confirm branch, operational date, posting period and teller/treasury assignment.")
bullet("Check opening cash or control balances according to local procedure.")
bullet("Review pending approvals, failed batches and unresolved exceptions.")
heading("Before every financial submission", 2)
bullet("Correct customer and account.")
bullet("Correct transaction type and debit/credit direction.")
bullet("Correct amount, currency, date and instrument/reference.")
bullet("Required evidence and approval are present.")
bullet("No duplicate transaction or batch exists.")
heading("End of day", 2)
bullet("Complete pending items or formally hand them over.")
bullet("Reconcile cash, transaction totals, batch totals and exceptions.")
bullet("Confirm all required authorisations or rejections have been recorded.")
bullet("Run and retain approved end-of-day reports.")
bullet("Use the Front Office End-of-Day function only when operational prerequisites are complete.")
bullet("Sign out and secure printed/exported information.")
callout("Local procedures prevail", "This guide explains application use. Your organisation's finance, cash, credit, privacy, approval and records-management policies remain mandatory.", "info")

page_break()
heading("15. Support and handover details", 1)
table(["Contact", "Details"], [
    ("Business system owner", "[Name / role / email / telephone]"),
    ("First-line support", "[Help desk / email / telephone / service hours]"),
    ("Technical escalation", "[Supplier / developer contact and escalation route]"),
    ("Security incident", "[Security contact and urgent reporting method]"),
    ("Production URL", "[URL]"),
    ("Training environment", "[URL, if available]"),
], [2.0, 4.72])
heading("Client acceptance", 2)
table(["Role", "Name", "Signature", "Date"], [
    ("Client system owner", "", "", ""),
    ("Operations representative", "", "", ""),
    ("Administrator", "", "", ""),
], [1.75, 1.8, 1.75, 1.42])
heading("Revision history", 2)
table(["Version", "Date", "Summary", "Owner"], [
    ("1.0", "September 2026", "Initial client user guide based on the delivered application build.", "[Owner]"),
], [0.8, 1.35, 3.35, 1.22])
callout("Before issue", "Replace all square-bracket placeholders, confirm enabled modules and have each process owner approve the relevant workflow section.", "warning")

doc.core_properties.title = "Swift Financial Client User Guide"
doc.core_properties.subject = "Client-facing user manual"
doc.core_properties.author = "Swift Financial"
doc.core_properties.comments = "Generated for client handover; replace bracketed placeholders before issue."
doc.save(DOCX)
print(DOCX)
